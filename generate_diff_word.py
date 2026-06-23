#!/usr/bin/env python3
"""Generate a Word doc mirroring what the web app shows for experienced track (diff modules).
All roles included. Uses proper Chinese fonts to avoid garbled characters."""

from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SKILL_ROOT = Path(__file__).resolve().parent
ARTICLES_DIR = SKILL_ROOT / "content" / "articles"

# ---- Font config ----
# macOS system Chinese fonts that render correctly in Word
CN_FONT = "STHeiti"        # 华文黑体 — good readability for body text
CN_FONT_HEADING = "STHeiti"  # Same for headings
EN_FONT = "Calibri"          # Western font for numbers/punctuation


def set_run_font(run, font_name=CN_FONT, size=11, bold=False, italic=False, color=None):
    """Set font for a run with proper Chinese font support."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set western font
    run.font.name = font_name
    # Set East-Asian font (critical for Chinese characters)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)


def add_p(doc, text, bold=False, size=11, italic=False, color=None):
    """Add a paragraph with proper Chinese font."""
    p = doc.add_paragraph()
    run = p.add_run(strip_emoji(text))
    set_run_font(run, CN_FONT, size, bold, italic, color)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_heading_styled(doc, text, level):
    """Add a heading with proper Chinese font."""
    h = doc.add_heading(strip_emoji(text), level=level)
    for run in h.runs:
        set_run_font(run, CN_FONT_HEADING, size={0:18, 1:15, 2:13, 3:12, 4:11}.get(level, 11), bold=True)
    return h


def parse_article(article_num: int) -> dict | None:
    """Parse an article markdown file into structured sections."""
    filepath = ARTICLES_DIR / f"article-{article_num:02d}.md"
    if not filepath.exists():
        return None
    text = filepath.read_text(encoding="utf-8")

    fm_match = re.match(r'^---\n(.*?)\n---\n', text, re.DOTALL)
    frontmatter = {}
    if fm_match:
        for line in fm_match.group(1).strip().split("\n"):
            if ":" in line:
                k, _, v = line.partition(":")
                frontmatter[k.strip()] = v.strip().strip('"').strip("[]").strip()
        body = text[fm_match.end():]
    else:
        body = text

    sections = {}
    current = None
    buf = []
    for line in body.split("\n"):
        m = re.match(r'^###\s+(.*)', line)
        if m:
            if current:
                sections[current] = "\n".join(buf).strip()
            current = m.group(1)
            buf = []
        else:
            buf.append(line)
    if current:
        sections[current] = "\n".join(buf).strip()

    return {
        "num": article_num,
        "chapter": frontmatter.get("chapter", ""),
        "title": frontmatter.get("title", ""),
        "change_level": frontmatter.get("change_level", ""),
        "sections": sections,
    }


def strip_emoji(text: str) -> str:
    """Remove emoji and other non-BMP characters that Word can't render."""
    import unicodedata
    result = []
    for ch in text:
        # Keep if not in Supplementary Multilingual Plane (emoji range) and not various symbols
        cp = ord(ch)
        if cp < 0x2100:  # below Miscellaneous Symbols
            result.append(ch)
        elif 0x2100 <= cp <= 0x27BF:  # Miscellaneous Symbols, Dingbats — strip emoji but keep basic symbols
            continue
        elif 0x1F000 <= cp <= 0x1FFFF:  # Emoticons, Symbols, etc.
            continue
        elif 0xFE00 <= cp <= 0xFE0F:  # Variation selectors
            continue
        else:
            result.append(ch)
    return "".join(result)


def clean_md(text: str) -> str:
    text = re.sub(r'^>\s?', '', text, flags=re.MULTILINE)
    text = text.replace("**", "")
    return text.strip()


def extract_all_roles(section_text: str) -> dict[str, str]:
    """Extract all role lines from the 👤 section."""
    roles = {}
    emoji_map = {"🏥": "PI", "📋": "CRA", "📝": "CRC", "🎓": "学生"}
    for line in section_text.split("\n"):
        line = line.strip()
        if line.startswith("- "):
            for emoji, name in emoji_map.items():
                if emoji in line:
                    cleaned = line.lstrip("- ").strip()
                    cleaned = re.sub(r'^[🏥📋📝🎓]\s*\*?\*?(PI|CRA|CRC|学生)\*?\*?[：:]\s*', '', cleaned)
                    if cleaned:
                        roles[name] = cleaned
                    break
    return roles


# ============================================================
# Module definitions (mirrors learn.py _get_diff_modules())
# ============================================================
MODULES = [
    {
        "title": "2026版 GCP 修订总览",
        "format": "💬",
        "description": "修订背景与立法逻辑：为什么从9章83条精简为6章54条，六章结构的功能逻辑，新增内容反映了哪些监管关切，术语变更的深层含义，与ICH E6(R3)的关系。",
        "duration": "~4 min",
        "articles": [],
    },
    {
        "title": "术语变更概览",
        "format": "💬",
        "description": "关键术语更新：试验参与者、主要研究者(PI)、服务供应商、伦理审查委员会。术语变更不仅是措辞调整，更反映了对临床试验参与关系和法律定位的重新理解。",
        "duration": "~2 min",
        "articles": [],
    },
    {
        "title": "核心理念升级",
        "format": "🔴",
        "description": "质量源于设计(QbD) + 风险相称的质量管理体系",
        "duration": "~6 min",
        "articles": [6, 40, 42],
    },
    {
        "title": "责任体系重构",
        "format": "🔴",
        "description": "PI不可授权事项 + 申办者质量管理体系",
        "duration": "~7 min",
        "articles": [19, 20, 32, 33],
    },
    {
        "title": "数据治理专章",
        "format": "🔴",
        "description": "电子源数据、元数据与稽查轨迹、计算机化系统验证、盲态数据管理",
        "duration": "~9 min",
        "articles": [51, 52, 53],
    },
    {
        "title": "知情同意 + 伦理审查变化",
        "format": "🟡",
        "description": "知情同意新要求 + 伦理审查委员会监督职能强化",
        "duration": "~5 min",
        "articles": [27, 14],
    },
    {
        "title": "新增专条",
        "format": "🔴",
        "description": "利益冲突回避 + 新技术新方法",
        "duration": "~5 min",
        "articles": [12, 13],
    },
    {
        "title": "修订要点回顾与路径选择",
        "format": "💬",
        "description": "总结2026版关键变化 + A/B/C分支选择说明",
        "duration": "~3 min",
        "articles": [],
    },
]


def build_doc():
    doc = Document()

    # ---- Page setup ----
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # ---- Set default paragraph font ----
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = CN_FONT
    style.paragraph_format.line_spacing = 1.3
    # Set east-asian font on Normal style
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = style.element.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)

    # ---- Title ----
    title = doc.add_heading('GCP 2026 培训 — 老手轨（差异速览）', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, CN_FONT_HEADING, 20, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('网页端用户所见内容 · 全部角色提示 · 供内部审核')
    set_run_font(r, CN_FONT, 13, color=(100, 100, 100))
    doc.add_paragraph()

    # ---- Module by module ----
    for mi, mod in enumerate(MODULES):
        if mi > 0:
            doc.add_page_break()
        add_heading_styled(doc, f'模块 {mi + 1}/8：{mod["title"]}', level=1)

        # Badge
        badge_text = mod["format"]
        if mod["format"] == "💬":
            badge_text += " 直接讲解"
        elif mod["format"] == "🔴":
            badge_text += " 深度讲授"
        elif mod["format"] == "🟡":
            badge_text += " 标准讲授"
        add_p(doc, f'【{badge_text}】', bold=True, size=10, color=(100, 100, 100))

        # Description + duration
        add_p(doc, f'{mod["description"]} · {mod["duration"]}', italic=True)

        if not mod["articles"]:
            # 💬 Module — web only shows description
            add_p(doc, '')
            add_p(doc, '（本模块为直接讲解，网页端仅展示以上描述文字，无条文内容。）', italic=True, size=10, color=(150, 150, 150))
            add_p(doc, '')
            add_p(doc, f'关于「{mod["title"]}」有疑问吗？', bold=True)
            add_p(doc, '[用户点击"继续 →"进入下一模块]', size=10, color=(128, 128, 128))
            continue

        # ---- 🔴/🟡 Module with articles ----
        for anum in mod["articles"]:
            art = parse_article(anum)
            if not art:
                add_p(doc, f'（第{anum}条文件未找到）', color=(255, 0, 0))
                continue

            sections = art["sections"]
            cl = art.get("change_level", "")

            # Article heading
            add_heading_styled(doc, f'第{anum}条 · {art["title"]}  {cl}', level=2)
            add_p(doc, f'所属章节：{art["chapter"]}', size=9, color=(128, 128, 128))

            # 💡 Scenario
            if "💡 先想一个场景" in sections:
                add_heading_styled(doc, '💡 先想一个场景', level=3)
                scenario_text = clean_md(sections["💡 先想一个场景"])
                if mod["format"] == "🔴":
                    add_p(doc, scenario_text, italic=True)
                else:
                    truncated = scenario_text[:120] + ("…" if len(scenario_text) > 120 else "")
                    add_p(doc, truncated, italic=True)

            # 📜 Original text
            if "📜 条文原文" in sections:
                add_heading_styled(doc, '📜 条文原文', level=3)
                add_p(doc, clean_md(sections["📜 条文原文"]), bold=True)

            # 🔑 Plain explanation
            if "🔑 这条在说什么" in sections:
                add_heading_styled(doc, '🔑 这条在说什么', level=3)
                add_p(doc, clean_md(sections["🔑 这条在说什么"]))

            # 🔄 Changes
            if "🔄 和旧版（2020）有什么不同" in sections:
                add_heading_styled(doc, '🔄 和旧版（2020）有什么不同', level=3)
                add_p(doc, clean_md(sections["🔄 和旧版（2020）有什么不同"]))

            # 👤 Role tips — ALL roles
            if "👤 对你的影响" in sections:
                add_heading_styled(doc, '👤 对你的影响（全部角色）', level=3)
                roles = extract_all_roles(sections["👤 对你的影响"])
                role_labels = [
                    ("PI", "🏥 主要研究者（PI）/ 医生"),
                    ("CRA", "📋 临床监查员（CRA）"),
                    ("CRC", "📝 临床研究协调员（CRC）"),
                    ("学生", "🎓 学生 / 初学者"),
                ]
                for rk, rl in role_labels:
                    if rk in roles:
                        p = doc.add_paragraph()
                        run = p.add_run(f'{rl}：')
                        set_run_font(run, CN_FONT, 11, bold=True)
                        run = p.add_run(roles[rk])
                        set_run_font(run, CN_FONT, 11)
                        p.paragraph_format.space_after = Pt(4)

            # ⚠️ Pitfalls
            if "⚠️ 容易踩的坑" in sections:
                if mod["format"] == "🔴":
                    add_heading_styled(doc, '⚠️ 容易踩的坑', level=3)
                else:
                    add_heading_styled(doc, '⚠️ 容易踩的坑（🟡默认折叠，用户点击展开）', level=3)
                add_p(doc, clean_md(sections["⚠️ 容易踩的坑"]), italic=True)

            # Separator between articles in same module
            if len(mod["articles"]) > 1 and anum != mod["articles"][-1]:
                doc.add_paragraph("—" * 40)

        # Module footer
        add_p(doc, '')
        add_p(doc, f'关于「{mod["title"]}」有疑问吗？', bold=True)
        add_p(doc, '[用户点击"继续 →"进入下一模块，也可点击"🤖 问 AI（本章）"]', size=10, color=(128, 128, 128))

    # ---- Appendix: 分支选择 ----
    doc.add_page_break()
    add_heading_styled(doc, '附录：分支选择页面', level=1)
    add_p(doc, '差异速览全部 8 个模块完成后，用户看到分支选择：')
    add_p(doc, '')
    add_p(doc, 'A. 只看差异，不做题', bold=True)
    add_p(doc, 'B. 仅变化题章节测验 — 只做变化题的章节测验，不做结业考试', bold=True)
    add_p(doc, 'C. 直接结业考试 — 不做章节测验，直接参加 30 题结业考试，通过发证', bold=True)
    add_p(doc, '')
    add_p(doc, '选择 B → 进入章节测验（change_only=1，只抽变化题）', size=10, color=(128, 128, 128))
    add_p(doc, '选择 C → 直接进入结业考试', size=10, color=(128, 128, 128))

    # ---- Save ----
    output = SKILL_ROOT / "GCP-2026-老手轨-用户所见内容-审核用.docx"
    doc.save(str(output))
    print(f"Done: {output}")
    return output


if __name__ == "__main__":
    build_doc()
